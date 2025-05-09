import os
import difflib
import html
from flask import current_app
from werkzeug.utils import secure_filename
import magic
import logging
from PIL import Image, ImageChops
from docx import Document
from pdfminer.high_level import extract_text

logger = logging.getLogger(__name__)

def save_file(file):
    filename = secure_filename(file.filename)
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    return filepath

def get_file_info(file_path):
    mime = magic.Magic(mime=True)
    file_type = mime.from_file(file_path)
    file_size = os.path.getsize(file_path)
    return {
        'name': os.path.basename(file_path),
        'type': file_type,
        'size': file_size,
        'origin': 'Uploaded from user device'
    }

def compare_files(file1, file2):
    mime1 = magic.from_file(file1, mime=True)
    mime2 = magic.from_file(file2, mime=True)
    if mime1 != mime2:
        return {"error": "Files must be of the same type for comparison."}, None
    if 'text' in mime1:
        return compare_text_files(file1, file2)
    elif 'image' in mime1:
        return compare_image_files(file1, file2)
    elif mime1 == 'application/pdf':
        return compare_pdf_files(file1, file2)
    elif mime1 == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return compare_docx_files(file1, file2)
    else:
        return {"error": "Unsupported file type."}, None

def compare_text_files(file1, file2):
    with open(file1, 'r', encoding='utf-8', errors='ignore') as f1, open(file2, 'r', encoding='utf-8', errors='ignore') as f2:
        text1 = f1.read()
        text2 = f2.read()

    matcher = difflib.SequenceMatcher(None, text1, text2)
    comparison = []
    differences = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            comparison.append({"status": "same", "file1": html.escape(text1[i1:i2]), "file2": html.escape(text2[j1:j2])})
        elif tag == 'replace':
            differences += 1
            comparison.append({"status": "different", "file1": html.escape(text1[i1:i2]), "file2": html.escape(text2[j1:j2])})
        elif tag == 'delete':
            differences += 1
            comparison.append({"status": "different", "file1": html.escape(text1[i1:i2]), "file2": ""})
        elif tag == 'insert':
            differences += 1
            comparison.append({"status": "different", "file1": "", "file2": html.escape(text2[j1:j2])})

    similarity = matcher.ratio() * 100
    detailed_analysis = generate_detailed_analysis(comparison)

    result_text = "\n".join([f"File1: {item['file1']} | File2: {item['file2']}" for item in comparison])

    output_folder = current_app.config.get('OUTPUT_FOLDER', 'output')
    os.makedirs(output_folder, exist_ok=True)
    output_path = os.path.join(output_folder, "comparison_result.txt")
    logger.debug(f"Saving comparison result to {output_path}")
    with open(output_path, 'w') as f:
        f.write(result_text)

    return {
        "similarity": round(similarity, 2),
        "differences": differences,
        "comparison": comparison,
        "detailedAnalysis": detailed_analysis
    }, output_path

def generate_detailed_analysis(comparison):
    analysis = []
    for item in comparison:
        if item['status'] == 'different':
            analysis.append(f"Difference found:\nFile1: {item['file1']}\nFile2: {item['file2']}")
    return "\n".join(analysis) if analysis else "Files are identical."

def compare_pdf_files(file1, file2):
    text1 = extract_text(file1)
    text2 = extract_text(file2)
    return compare_text_files(file1, file2)

def compare_docx_files(file1, file2):
    doc1 = Document(file1)
    doc2 = Document(file2)
    text1 = '\n'.join([p.text for p in doc1.paragraphs])
    text2 = '\n'.join([p.text for p in doc2.paragraphs])
    return compare_text_files(file1, file2)

def compare_image_files(file1, file2):
    try:
        img1 = Image.open(file1)
        img2 = Image.open(file2)
        if img1.size != img2.size:
            return {"result": "Images have different dimensions.", "similarity": 0, "differences": 1}, None
        diff = ImageChops.difference(img1, img2)
        if diff.getbbox() is None:
            result = {"result": "Images are identical.", "similarity": 100, "differences": 0}
            output_path = None
        else:
            output_folder = current_app.config.get('OUTPUT_FOLDER', 'output')
            os.makedirs(output_folder, exist_ok=True)
            output_path = os.path.join(output_folder, 'diff_image.png')
            diff.save(output_path)
            result = {"result": "Images are different. Difference image saved.", "similarity": 0, "differences": 1}
        return result, output_path
    except Exception as e:
        logger.error(f"Error comparing images: {e}")
        return {"error": f"An error occurred while processing images: {e}"}, None